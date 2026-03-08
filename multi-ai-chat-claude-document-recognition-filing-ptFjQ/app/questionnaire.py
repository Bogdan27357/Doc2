"""Generate employee questionnaire (личная карточка работника) as DOCX."""

import io

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _set_cell(cell, text: str, bold: bool = False, size: int = 10):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def generate_questionnaire(fields: dict) -> io.BytesIO:
    """Generate a DOCX questionnaire from extracted document fields.

    Args:
        fields: dict with keys like surname, name, patronymic, birth_date,
                gender, birth_place, series, number, issue_date, issued_by,
                department_code, snils_number, inn_number, institution,
                specialty, qualification, etc.

    Returns:
        BytesIO buffer with the DOCX file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ЛИЧНАЯ КАРТОЧКА РАБОТНИКА")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("(Унифицированная форма № Т-2)")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph()  # spacing

    # Helper to get field value
    def f(key: str) -> str:
        return fields.get(key, "") or ""

    full_name = f"{f('surname')} {f('name')} {f('patronymic')}".strip()

    # ---- Section 1: General Info ----
    h = doc.add_paragraph()
    run = h.add_run("I. ОБЩИЕ СВЕДЕНИЯ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    table1 = doc.add_table(rows=10, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = "Table Grid"

    rows_data = [
        ("Фамилия", f("surname")),
        ("Имя", f("name")),
        ("Отчество", f("patronymic")),
        ("Дата рождения", f("birth_date")),
        ("Место рождения", f("birth_place")),
        ("Пол", f("gender")),
        ("Гражданство", "Российская Федерация"),
        ("Документ, удостоверяющий личность",
         f"Паспорт серия {f('series')} № {f('number')}" if f("series") else ""),
        ("Дата выдачи паспорта", f("issue_date")),
        ("Кем выдан", f("issued_by")),
    ]

    for i, (label, value) in enumerate(rows_data):
        _set_cell(table1.rows[i].cells[0], label, bold=True)
        _set_cell(table1.rows[i].cells[1], value)

    # Set column widths
    for row in table1.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()  # spacing

    # ---- Section 2: IDs ----
    h = doc.add_paragraph()
    run = h.add_run("II. ИДЕНТИФИКАЦИОННЫЕ НОМЕРА")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    table2 = doc.add_table(rows=3, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Table Grid"

    ids_data = [
        ("ИНН", f("inn_number")),
        ("СНИЛС", f("snils_number")),
        ("Код подразделения (паспорт)", f("department_code")),
    ]

    for i, (label, value) in enumerate(ids_data):
        _set_cell(table2.rows[i].cells[0], label, bold=True)
        _set_cell(table2.rows[i].cells[1], value)

    for row in table2.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()

    # ---- Section 3: Education ----
    h = doc.add_paragraph()
    run = h.add_run("III. ОБРАЗОВАНИЕ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    table3 = doc.add_table(rows=5, cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = "Table Grid"

    edu_data = [
        ("Учебное заведение", f("institution")),
        ("Специальность", f("specialty")),
        ("Квалификация", f("qualification")),
        ("Документ об образовании",
         f"Серия {f('diploma_series')} № {f('diploma_number')}"
         if f("diploma_series") else ""),
        ("Дата выдачи диплома", f("diploma_issue_date") or f("edu_issue_date")),
    ]

    for i, (label, value) in enumerate(edu_data):
        _set_cell(table3.rows[i].cells[0], label, bold=True)
        _set_cell(table3.rows[i].cells[1], value)

    for row in table3.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()

    # ---- Signatures ----
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig.add_run(f"Работник: __________________ / {full_name} /")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig2.add_run("Дата заполнения: «____» ______________ 20___ г.")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
